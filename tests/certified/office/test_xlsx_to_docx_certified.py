from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.core.settings import settings
from app.main import app


def test_xlsx_to_docx_conversion_creates_valid_docx(tmp_path: Path):
    client = TestClient(app)

    xlsx_path = Path("tests/assets/regression/sample.xlsx")
    assert xlsx_path.exists(), "Sample XLSX is missing"

    resp = client.post(
        "/convert",
        data={"target_format": "docx", "operation": "xlsx-to-docx"},
        files={
            "file": (
                xlsx_path.name,
                xlsx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )

    assert resp.status_code == 201, resp.text
    payload = resp.json()
    assert payload.get("status") == "success"
    download_path = payload.get("download_path")
    assert download_path, payload
    assert download_path.startswith("/download/")
    relative_parts = Path(download_path.removeprefix("/download/")).parts
    assert len(relative_parts) == 2, f"Unexpected download path shape: {download_path}"
    conversion_id, filename = relative_parts
    assert filename.endswith(".docx")

    local_path = settings.OUTPUT_DIR / conversion_id / filename
    download_resp = client.get(download_path)
    assert download_resp.status_code == 200, download_resp.text
    assert download_resp.content, "Downloaded content is empty"

    assert local_path.exists(), f"Expected output DOCX not found: {local_path}"
    assert local_path.stat().st_size > 0, "Output DOCX is empty"
    assert local_path.suffix.lower() == ".docx"

    # Content-integrity: re-open output with python-docx and verify the sheet
    # heading and the cell data from the source XLSX actually made it in.
    doc = DocxDocument(str(local_path))
    texts = [p.text for p in doc.paragraphs if p.text]
    assert texts, "DOCX output has no paragraphs"

    joined = "\n".join(texts)
    assert "Sheet1" in joined, f"Expected sheet heading in DOCX, got: {joined[:500]}"

    tables = doc.tables
    assert tables, "DOCX output has no table"

    table_cells = [cell.text for row in tables[0].rows for cell in row.cells]
    assert "Sample" in table_cells, f"Expected cell data from XLSX in DOCX table: {table_cells}"
    assert "123" in table_cells, f"Expected numeric cell from XLSX in DOCX table: {table_cells}"
    assert len(list((settings.OUTPUT_DIR / conversion_id).glob("*"))) == 1


def test_xlsx_to_docx_output_is_real_ooxml_docx(tmp_path: Path):
    """Verify the output starts with the OOXML ZIP magic bytes (PK), proving it
    is a real DOCX container rather than a plain-text/fake file."""
    client = TestClient(app)

    xlsx_path = Path("tests/assets/regression/sample.xlsx")
    resp = client.post(
        "/convert",
        data={"target_format": "docx", "operation": "xlsx-to-docx"},
        files={
            "file": (
                xlsx_path.name,
                xlsx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert resp.status_code == 201, resp.text
    payload = resp.json()
    conversion_id, filename = Path(payload["download_path"].removeprefix("/download/")).parts
    local_path = settings.OUTPUT_DIR / conversion_id / filename

    with local_path.open("rb") as handle:
        header = handle.read(4)
    assert header == b"PK\x03\x04", f"Output is not a real DOCX (ZIP) file: {header!r}"
