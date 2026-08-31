"""
Project : Converigo
Author  : Pico Lala & ChatGPT
Version : 3.0.1

DOCX -> XLSX Plugin

Converts Word documents into Excel spreadsheets. Source detection is
content-based (magic bytes) rather than extension-based:

- ZIP/OOXML magic (PK) -> treated as DOCX and real document content is extracted,
  which also covers `.doc` files that are actually DOCX renamed.
- OLE2/CFB magic (d0cf11e0a1b11ae1) -> legacy binary `.doc`; not supported,
  raises an explicit error instead of emitting an empty XLSX.
- Anything else -> raises an explicit error (no silent fake pass).
"""

from pathlib import Path

from app.plugins.base import ConverterPlugin


class DOCXToXLSXPlugin(ConverterPlugin):
    slug = "docx-to-xlsx"
    name = "DOCX to XLSX"
    description = "Convert Word documents into Excel spreadsheets."
    category = "document"
    engine = "document"
    icon = "📄"

    source_formats = ["docx", "doc", "word"]
    target_formats = ["xlsx", "xls", "spreadsheet"]

    goal = "document"
    use_case = "Best for extracting Word document text and tables into Excel spreadsheets."
    priority = 80
    quality = 85
    compatibility = 80
    estimated_saving = 10
    badge = "Office Conversion"
    seo_title = "DOCX to XLSX Converter | Converigo"
    seo_description = "Convert Word documents into Excel spreadsheets quickly and easily."

    # DOCX (OOXML) files are ZIP archives and start with the "PK" signature.
    _DOCX_PK_MAGIC = b"PK\x03\x04"
    # Legacy .doc files are OLE2 / Compound File Binary containers.
    _DOC_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

    @classmethod
    def _detect_container(cls, source_path: Path) -> str:
        """Return the container type detected from file content.

        Returns "docx" for ZIP/OOXML files, "doc" for OLE2/CFB files, or
        "unknown" for anything else (empty file, RTF, plain text, ...).
        """
        try:
            with source_path.open("rb") as fh:
                header = fh.read(8)
        except OSError as exc:
            raise RuntimeError(f"Could not read the source document: {exc}") from exc

        if header.startswith(cls._DOCX_PK_MAGIC):
            return "docx"
        if header.startswith(cls._DOC_OLE2_MAGIC):
            return "doc"
        return "unknown"

    @staticmethod
    def _paragraph_text(paragraph) -> str:
        """Return the stripped text of a document paragraph."""
        return paragraph.text.strip()

    @staticmethod
    def _row_values(row) -> list[str]:
        """Return the cell texts of a table row as a flat list of strings."""
        return [cell.text.strip() for cell in row.cells]

    async def convert(
        self,
        source_path: Path,
        target_format: str,
        output_dir: Path | None = None,
        temp_dir: Path | None = None,
    ) -> Path:
        if not self.supports(source_path.suffix, target_format):
            raise RuntimeError("DOCXToXLSXPlugin only supports DOCX/DOC -> XLSX.")

        container = self._detect_container(source_path)

        if container == "doc":
            raise RuntimeError(
                "Legacy .doc format (OLE2 compound document) is not supported by "
                "this converter. Please save your document as .docx and try again."
            )
        if container == "unknown":
            raise RuntimeError(
                "The uploaded file is not a valid DOCX document. Please save it "
                "as a valid .docx file and try again."
            )

        try:
            from docx import Document

            document = Document(str(source_path))
        except Exception as exc:
            raise RuntimeError(
                "The DOCX document could not be parsed. Please save it as a valid "
                ".docx file and try again."
            ) from exc

        import openpyxl

        workbook = openpyxl.Workbook()
        workbook.remove(workbook.active)

        # Sheet "Content": one row per non-empty paragraph.
        content_sheet = workbook.create_sheet("Content", 0)
        content_sheet.append(["Paragraph Text"])
        for paragraph in document.paragraphs:
            text = self._paragraph_text(paragraph)
            if text:
                content_sheet.append([text])

        # One sheet per native Word table (cell layout preserved).
        for index, table in enumerate(document.tables, start=1):
            sheet = workbook.create_sheet(f"Table {index}")
            for row in table.rows:
                sheet.append(self._row_values(row))

        from app.core.settings import settings

        working_dir = temp_dir or output_dir or (settings.OUTPUT_DIR / "document")
        working_dir.mkdir(parents=True, exist_ok=True)

        output_path = working_dir / f"{source_path.stem}.xlsx"
        workbook.save(str(output_path))
        workbook.close()

        if not output_path.exists() or output_path.stat().st_size == 0:
            raise RuntimeError("DOCX/DOC to XLSX conversion did not produce output.")

        return output_path
