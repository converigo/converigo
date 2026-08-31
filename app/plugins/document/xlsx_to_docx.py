"""
Project : Converigo
Author  : Pico Lala & ChatGPT
Version : 3.0.1

XLSX -> DOCX Plugin

Converts Excel spreadsheets into Word documents. Source detection is
content-based (magic bytes) rather than extension-based:

- ZIP/OOXML magic (PK) -> treated as XLSX and real cell data is extracted,
  which also covers `.xls` files that are actually XLSX renamed.
- OLE2/CFB magic (d0cf11e0a1b11ae1) -> legacy binary `.xls`; not supported,
  raises an explicit error instead of emitting a content-less DOCX.
- Anything else -> raises an explicit error (no silent fake pass).
"""

from pathlib import Path

from app.plugins.base import ConverterPlugin


class XLSXToDOCXPlugin(ConverterPlugin):
    slug = "xlsx-to-docx"
    name = "XLSX to DOCX"
    description = "Convert Excel spreadsheets into Word documents."
    category = "document"
    engine = "document"
    icon = "📄"

    source_formats = ["xlsx", "xls", "spreadsheet"]
    target_formats = ["docx", "doc", "word"]

    goal = "document"
    use_case = "Best for turning spreadsheet data into editable Word documents."
    priority = 80
    quality = 85
    compatibility = 80
    estimated_saving = 10
    badge = "Office Conversion"
    seo_title = "XLSX to DOCX Converter | Converigo"
    seo_description = "Convert Excel spreadsheets into Word documents quickly and easily."

    # XLSX (OOXML) files are ZIP archives and start with the "PK" signature.
    _XLSX_PK_MAGIC = b"PK\x03\x04"
    # Legacy .xls files are OLE2 / Compound File Binary containers.
    _XLS_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

    @classmethod
    def _detect_container(cls, source_path: Path) -> str:
        """Return the container type detected from file content.

        Returns "xlsx" for ZIP/OOXML files, "xls" for OLE2/CFB files, or
        "unknown" for anything else (empty file, RTF, plain text, ...).
        """
        try:
            with source_path.open("rb") as fh:
                header = fh.read(8)
        except OSError as exc:
            raise RuntimeError(f"Could not read the source spreadsheet: {exc}") from exc

        if header.startswith(cls._XLSX_PK_MAGIC):
            return "xlsx"
        if header.startswith(cls._XLS_OLE2_MAGIC):
            return "xls"
        return "unknown"

    @staticmethod
    def _cell_text(value) -> str:
        """Render a cell value as plain text for a Word table cell."""
        if value is None:
            return ""
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value)

    async def convert(
        self,
        source_path: Path,
        target_format: str,
        output_dir: Path | None = None,
        temp_dir: Path | None = None,
    ) -> Path:
        if not self.supports(source_path.suffix, target_format):
            raise RuntimeError("XLSXToDOCXPlugin only supports XLSX/XLS -> DOCX.")

        container = self._detect_container(source_path)

        if container == "xls":
            raise RuntimeError(
                "Legacy .xls format (OLE2 compound document) is not supported by "
                "this converter. Please save your spreadsheet as .xlsx and try again."
            )
        if container == "unknown":
            raise RuntimeError(
                "The uploaded file is not a valid XLSX spreadsheet. Please save it "
                "as a valid .xlsx file and try again."
            )

        try:
            import openpyxl

            workbook = openpyxl.load_workbook(
                str(source_path),
                data_only=True,
                read_only=True,
            )
        except Exception as exc:
            raise RuntimeError(
                "The XLSX spreadsheet could not be parsed. Please save it as a valid "
                ".xlsx file and try again."
            ) from exc

        from docx import Document

        document = Document()

        try:
            for worksheet in workbook.worksheets:
                document.add_heading(worksheet.title, level=1)

                rows = [row for row in worksheet.iter_rows(values_only=True) if any(v is not None for v in row)]
                if not rows:
                    continue

                cols = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=cols)
                for r, row in enumerate(rows):
                    for c in range(cols):
                        value = row[c] if c < len(row) else None
                        table.cell(r, c).text = self._cell_text(value)
                try:
                    table.style = "Table Grid"
                except Exception:
                    # Style is cosmetic only; a valid table without it is still fine.
                    pass
        finally:
            workbook.close()

        from app.core.settings import settings

        working_dir = temp_dir or output_dir or (settings.OUTPUT_DIR / "document")
        working_dir.mkdir(parents=True, exist_ok=True)

        output_path = working_dir / f"{source_path.stem}.docx"
        document.save(str(output_path))

        if not output_path.exists() or output_path.stat().st_size == 0:
            raise RuntimeError("XLSX/XLS to DOCX conversion did not produce output.")

        return output_path
